# -*- coding: utf-8 -*-
"""
individual_report.py
=====================
Word report for the personal PYROX assessment (app_persoonlijk.py /
individual_engine.py).

Deliberately a SEPARATE module from report_generator.py (which serves
the population-level apps) rather than a wrapper around it:
IndividualAssessment's shape -- minutes+phase spanning race AND
post-finish, a personal ensemble instead of a sampled population,
representative_trajectories(), zone_episode() -- doesn't map onto the
population report's per-level hestia_result dicts. Forcing it to would
either require computing data this app doesn't have, or silently
reintroduce population-level assumptions into a report about one
person.

Reuses report_generator.py's established building blocks wherever the
underlying data shapes genuinely match, rather than re-implementing
document styling a second time:
    _add_title_section, _add_heading, _add_caption,
    _disable_first_row_style, _weather_chart, _t_rect_co_reserve_scatter,
    dose_evolution_chart
dose_evolution_chart() in particular needs no adaptation at all --
representative_trajectories() calls the exact same hestia_bridge
selection function the population apps use, so the shape already
matches.

PRIVACY: this function takes only already-computed, in-memory objects
(PersonalInputs, EventScenario, IndividualAssessment) and returns bytes
directly. No network access, no file written to disk, no telemetry. The
resulting bytes never leave the caller's control until Streamlit's own
download_button hands them to the browser that requested them.

LANGUAGE: English, matching every other Word/Excel report this suite
already produces (see report_generator.py's own paragraphs). The
Streamlit UI is Dutch, but generated documents have been English
throughout this project -- said explicitly here rather than assumed
silently, since it's an easy thing to want the other way.
"""

from __future__ import annotations

import io

import pandas as pd
from docx.shared import Pt, Cm, RGBColor

from individual_engine import (
    PersonalInputs, EventScenario, IndividualAssessment,
    representative_trajectories, dose_scatter_points, zone_episode,
    assessment_caveats,
)
from hestia_model import PF_DUUR_MIN
import report_generator as rg


#: hestia_bridge's own trace labels (English, used identically by the
#: population apps) -- kept in English here for the Word report, unlike
#: app_persoonlijk.py's Dutch UI translation of the same labels.
_LABEL_EN = {
    "Population median": "Ensemble median",
    "Lowest risk (dose=0)": "Lowest risk (dose=0)",
    "Median non-zero dose": "Median non-zero dose",
    "Highest dose": "Highest dose",
    "Coolest participant": "Coolest run",
    "Median participant": "Median run",
    "Hottest participant": "Hottest run",
}


def _zone_episode_text(label: str, ep: dict) -> str:
    """English counterpart of app_persoonlijk.py's Dutch zone_episode
    explanation -- same three patterns, same underlying facts, worded
    for a document rather than a Streamlit caption. Keep both in sync
    if the classification logic in zone_episode() ever changes."""
    if ep["entered_only_postfinish"]:
        return (
            f"{label}: entered the danger zone only AFTER finishing -- never during the "
            f"race itself. T_rect kept rising post-finish until it crossed the threshold "
            f"(a delayed-heat-storage effect); stopping was not, by itself, sufficient "
            f"protection here."
        )
    if ep["exited_during_race"]:
        return (
            f"{label}: entered the zone and left it again WHILE STILL RUNNING. The body "
            f"pulled back under load -- the most reassuring of the three patterns."
        )
    if ep["in_zone_at_finish"]:
        return (
            f"{label}: was still inside the danger zone AT the finish. Any recovery "
            f"afterward is almost certainly because exertion stopped (cardiac demand drops "
            f"sharply at rest), not because the thermal danger had resolved."
        )
    return f"{label}: entered the zone."


def generate_individual_report_docx(
    inputs: PersonalInputs,
    scenario: EventScenario,
    result: IndividualAssessment,
    app_build: str = "individual-1",
) -> bytes:
    """A Word report for one person's own PYROX assessment.

    Same scope discipline as report_generator.py's population reports:
    describes what the model found, not what to do about it. Operates
    entirely on the already-computed `result` -- never re-runs the
    simulation, never makes a network call.
    """
    doc = rg.Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)

    generated_at = pd.Timestamp.now(tz="UTC")
    rg._add_title_section(
        doc, result.city_name, result.meteo["time"][0], generated_at, app_build,
        report_title="PYROX \u2014 Personal Heat-Risk Assessment",
        scope_text=(
            "SCOPE: This report presents model outputs and findings only, for one "
            "person's own entered details and one event. It is not medical advice and "
            "does not recommend whether to take part, how to pace, or what to do about "
            "the result -- consult a medical professional for guidance specific to your "
            "health. Where two figures in this report disagree, an explanation of WHY is "
            "included where relevant; that is a methodological note, not advice."
        ),
    )

    priv = doc.add_paragraph()
    priv_run = priv.add_run(
        "\U0001F512 PRIVACY: this report was generated entirely on the device that ran the "
        "assessment. The only network calls anywhere in this pipeline are geocoding the "
        "event location and fetching weather for it -- personal details (height, weight, "
        "age, pace, ...) and every result derived from them never left this machine."
    )
    priv_run.italic = True
    priv_run.font.size = Pt(9)
    priv_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()

    # --- Personal profile ---------------------------------------------
    rg._add_heading(doc, "Personal profile", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    rg._disable_first_row_style(table)

    def _row(label: str, value) -> None:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = str(value)

    _row("Height", f"{inputs.height_m:.2f} m")
    _row("Weight", f"{inputs.weight_kg:.1f} kg")
    _row("Age", f"{inputs.age}")
    _row("Gender", inputs.gender)
    _row("Expected pace", f"{inputs.expected_pace_min_per_km:.1f} min/km")
    _row("NSAID use", "Yes" if inputs.nsaid_use else "No")
    _row("Drinking habit",
         "Drinks at most stations" if inputs.drinks_readily else "Waits until clearly thirsty")
    _row("Heat acclimatized",
         "Yes (trained/lived in heat recently)" if inputs.heat_acclimatized else "No")
    _row("VO2max",
         f"{inputs.known_vo2max:.1f} mL/kg/min (entered)" if inputs.known_vo2max
         else "estimated from age/gender (not entered)")
    _row("Body fat %",
         f"{inputs.known_body_fat_pct:.1f}% (entered)" if inputs.known_body_fat_pct
         else "estimated from gender (not entered)")
    rg._add_caption(
        doc, "Fields marked 'estimated' were not entered and default to an age/gender-"
             "conditioned population mean, not a measurement of this specific person."
    )
    doc.add_paragraph()

    # --- Event ----------------------------------------------------------
    rg._add_heading(doc, "Event", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"{result.city_name} \u2014 "
        f"{pd.Timestamp(scenario.start_local).strftime('%A %d %B %Y, %H:%M')}"
    ).bold = True
    doc.add_paragraph(
        f"Duration: {scenario.duration_minutes:.0f} minutes. "
        f"{'Historical (observed weather)' if scenario.use_historical else 'Forecast'}. "
        f"Ensemble size: {result.n_ensemble} personal simulation runs "
        f"(same physiology engine as PYROX's population apps, run for this one person's "
        f"own entered details instead of a sampled population)."
    )

    # --- Weather (reuse _weather_chart via a minimal reconstructed frame) ---
    meteo = result.meteo
    weather_df_mini = pd.DataFrame(
        {"T_air_urban": meteo["t_air"], "WBGT": meteo["wbgt"],
         "UTCI": meteo["utci"], "MRT": meteo["mrt"]},
        index=pd.DatetimeIndex(meteo["time"]),
    )
    exp_start = meteo["time"][0]
    finish_ts = meteo["time"][-1]
    post_finish_end = finish_ts + pd.Timedelta(minutes=PF_DUUR_MIN)
    png, _plotted = rg._weather_chart(weather_df_mini, exp_start, finish_ts, post_finish_end)
    if png:
        rg._add_heading(doc, "Weather conditions during the race", level=1)
        doc.add_picture(png, width=Cm(15.5))
        rg._add_caption(
            doc, "Same colour mapping as the live app's weather chart. The shaded band "
                 "after the finish marks HESTIA's fixed post-finish simulation window -- "
                 "no new weather data applies there."
        )

    # --- EHS estimate -----------------------------------------------------
    rg._add_heading(doc, "Estimated EHS likelihood", level=1)
    ehs = result.ehs_interval
    if "error" not in ehs:
        pct = ehs["point_per_1000"] / 10.0
        pct_lo = ehs["lo_per_1000"] / 10.0
        pct_hi = ehs["hi_per_1000"] / 10.0
        p = doc.add_paragraph()
        p.add_run(
            f"\u2248{pct:.2f}% ({pct_lo:.2f}\u2013{pct_hi:.2f}%, "
            f"95% sampling + anchor interval)"
        ).bold = True
        doc.add_paragraph(
            "Based on epidemiology of comparable runners (Falmouth Road Race data), "
            "applied to this person's simulated dose -- not a probability independently "
            "validated for this individual."
        )
    p = doc.add_paragraph()
    p.add_run(
        f"Conjunctive criterion reached (T_rect\u226540.5\u00b0C AND CO_reserve\u22640): "
        f"{result.conjunction_fraction:.0%} of ensemble runs"
    ).bold = True

    p = doc.add_paragraph()
    p.add_run(
        f"EHE \u2014 exertional heat exhaustion (T_rect>39.5\u00b0C AND CO_reserve<0, "
        f"during exertion): {result.ehe_fraction:.0%} of ensemble runs; "
        f"mean dose {result.ehe_dose_mean:.2f}, median among affected runs {result.ehe_dose_among_hits:.2f}"
    ).bold = True
    p = doc.add_paragraph()
    p.add_run(
        f"EAC \u2014 exercise-associated collapse (CO_reserve<0 post-finish, no "
        f"temperature condition): {result.eac_fraction:.0%} of ensemble runs; "
        f"mean dose {result.eac_dose_mean:.2f}, median among affected runs {result.eac_dose_among_hits:.2f}"
    ).bold = True

    rg._add_heading(doc, "What these three criteria mean", level=2)
    for term, body in [
        ("EHS \u2014 Exertional Heat Stroke",
         "Clinically: CNS dysfunction together with a core temperature above "
         "40.5\u00b0C (Roberts 2010; ACSM 2023; NATA). This model cannot simulate "
         "neurological status, so it substitutes cardiovascular decompensation "
         "(CO_reserve\u22640) for the CNS criterion. That substitution is "
         "conservative: cerebral hypoperfusion has been measured at 40\u00b0C core "
         "temperature with cardiac output still intact (Nybo & Nielsen 2001), so "
         "CNS dysfunction can occur before CO_reserve reaches zero. Systematic "
         "under-detection is absorbed by calibrating the intercept against "
         "observed EHS counts (Breslow et al. 2021, Boston Marathon)."),
        ("EHE \u2014 Exertional Heat Exhaustion",
         "Clinically: inability to continue, core temperature typically "
         "38.5\u201340\u00b0C, WITHOUT the CNS dysfunction that defines EHS "
         "(ACSM 2023). Modelled here as T_rect>39.5\u00b0C and CO_reserve<0 at the "
         "same timestep, during exertion. Checked against this model's own output: "
         "the state does NOT then progress to 40.5\u00b0C \u2014 temperature "
         "plateaus at a genuine thermal steady state while CO_reserve keeps "
         "eroding through dehydration, with the metabolic rate essentially "
         "unchanged. It therefore marks LOST CONTROL MARGIN rather than impending "
         "heat stroke: temperature holds only because production and loss happen "
         "to balance, while the capacity to absorb any further disturbance "
         "disappears. The dose (deficit integrated over time) carries this signal; "
         "the yes/no flag does not."),
        ("EAC \u2014 Exercise-Associated Collapse",
         "Clinically: a conscious athlete unable to stand or walk unaided after "
         "an endurance event, caused by postural hypotension when the muscle pump "
         "stops at the finish line while skin vessels stay dilated (Asplund & "
         "O'Connor 2011; Roberts 2007). It is cardiovascular, not thermal, so no "
         "temperature threshold is applied \u2014 that omission is deliberate, not "
         "an oversight. Collapse DURING a race points to a different and more "
         "serious cause. Reference incidence for future calibration: 1.53 per 1000 "
         "runners (Gothenburg Half Marathon); EAC accounts for 59\u201385% of "
         "finish-line medical-tent visits."),
        ("CO_reserve \u2014 cardiac output reserve",
         "The share of maximum cardiac output not currently claimed by the demands "
         "of exercise and thermoregulation combined (Lloyd et al. 2022). Zero means "
         "the actuator is saturated: no further increase in cooling capacity is "
         "available, though cooling already achieved continues. Negative values "
         "indicate demand exceeding the heat-reduced maximum."),
        ("Dose",
         "The deficit integrated over time (L/min \u00d7 minutes) while a criterion "
         "holds \u2014 weighting both how deep and how long, rather than treating "
         "every qualifying moment alike."),
        ("Conjunctive / simultaneous",
         "All criteria here require both conditions at the SAME timestep. A "
         "temperature peak at 11:00 and a reserve trough at 13:00 do not count. "
         "This is the distinction from additive indices such as WBGT, and from "
         "models that read each variable's extreme independently."),
    ]:
        para = doc.add_paragraph()
        para.add_run(term + ". ").bold = True
        para.add_run(body)

    for c in assessment_caveats(result):
        cp = doc.add_paragraph()
        cr = cp.add_run("\u26a0 " + c)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x7F, 0x1D, 0x1D)
    doc.add_paragraph()

    # --- T_rect / CO_reserve over time (reuse dose_evolution_chart) -------
    reps = representative_trajectories(result.all_traces)
    png = rg.dose_evolution_chart(reps, level_label="")
    if png:
        rg._add_heading(doc, "How T_rect and CO_reserve evolved", level=1)
        doc.add_picture(png, width=Cm(15.5))
        rg._add_caption(
            doc, "Ensemble median shown as a thick dashed line; the coloured lines are "
                 "individually-picked representative runs from this personal ensemble, "
                 "not other people."
        )
        for tr in reps:
            label = _LABEL_EN.get(tr["label"], tr["label"])
            ep = zone_episode(tr)
            if ep is not None:
                doc.add_paragraph(_zone_episode_text(label, ep))

    # --- T_rect vs CO_reserve scatter (reuse _t_rect_co_reserve_scatter) --
    cloud = dose_scatter_points(result.all_traces)
    pairs = list(zip(cloud["t_rect"].tolist(), cloud["co_reserve"].tolist()))
    png = rg._t_rect_co_reserve_scatter(pairs, level_label="")
    if png:
        rg._add_heading(doc, "T_rect vs CO_reserve", level=1)
        doc.add_picture(png, width=Cm(13.5))
        rg._add_caption(
            doc, "One point per (ensemble member, timestep), race and post-finish combined."
        )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        "This report presents model outputs and findings only -- it does not include "
        "recommendations about whether to run, how to pace, or when to seek medical "
        "attention. Consult a medical professional for guidance specific to your health."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
