# -*- coding: utf-8 -*-
"""
Report generator — findings only, deliberately no recommendations
=====================================================================
Produces a Word report from a run's data across all layers (WBGT flags,
PYROX multi-day reserve, HESTIA acute capacity). Explicitly scoped to
DESCRIBE what the models found, not to PRESCRIBE what to do about it --
no staffing numbers, no "we advise", no start-time suggestions. That
judgment belongs to the organising/medical team, not to a model output.

The one exception, and it is a deliberate one: explaining WHY two numbers
differ (e.g. "the flag and the reserve disagree because they operate at
different timescales") is METHODOLOGICAL guidance on how to read the
data, not an operational recommendation about what to do. That
distinction is maintained throughout -- if a sentence would tell someone
what action to take, it doesn't belong in this report.
"""

from __future__ import annotations

__BUILD__ = "2026-08-14a"

import io
from datetime import datetime

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")  # headless -- no display available server-side
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# =============================================================================
# Small formatting helpers
# =============================================================================
def dose_positive_count(hestia_result: dict) -> tuple[int, int]:
    """(n_participants_with_dose_above_zero, n_total) from a HESTIA result
    dict. Distinct PARTICIPANTS, not timestep-points -- the T_rect/CO_reserve
    scatter's own caption already warns that its point count is per
    timestep, which over-counts a participant who lingers in the danger
    quadrant for several consecutive steps. This is the participant-level
    count the headline EHS estimate is actually an average over."""
    doses = hestia_result.get("cumulative_doses_all", []) if hestia_result else []
    n_total = len(doses)
    n_positive = sum(1 for d in doses if d is not None and d > 0)
    return n_positive, n_total


#: Below this many participants with a non-zero dose, the mean-of-logistic
#: headline estimate is disproportionately shaped by one or two individuals
#: rather than reflecting a broad pattern -- see the docstring on
#: dose_response_ehs_probability's caller in hestia_bridge.py for why: the
#: curve saturates near 1.0 for a high dose, so a single such participant
#: can move the population mean by several points per 1000 on its own,
#: regardless of how many simulations were run in total.
DOSE_POSITIVE_WARNING_THRESHOLD = 10


def dose_positive_warning_text(n_positive: int, n_total: int,
                               broad_screen_pct: float | None = None,
                               mean_t_air_c: float | None = None) -> str | None:
    """Human-readable note if the EHS estimate needs context beyond the
    bare number, or None if it can stand on its own. Two distinct cases,
    both real and both worth surfacing -- shared wording between the live
    app and the Word report:

    n_positive == 0 : NONE of the simulated participants ever reached a
    non-zero dose. [2026-08-13 patch] For this case the headline figure
    is no longer the logistic curve's own fixed global intercept -- it is
    Falmouth's own real-world temperature-only rate (falmouth_ehs_per_1000)
    at today's actual race-window temperature, which now correctly scales
    down at cooler temperatures instead of staying pinned at a constant
    ~1.7 per 1000 regardless of how mild the day was. See
    _dose_response_pct_patched's docstring in hestia_bridge.py for why a
    full re-fit of the dose-slope itself was tried and NOT shipped (dose
    and temperature turned out to be confounded at the reference
    calibration scenario, leaving no way to identify a dose effect
    separately from a temperature effect with the available data).

    0 < n_positive < DOSE_POSITIVE_WARNING_THRESHOLD : the estimate is
    real but rests on very few participants, so it may reflect a couple
    of individuals rather than a broad pattern (see below). This case is
    UNCHANGED by the patch -- dose>0 behaviour still uses the original
    logistic curve.
    """
    if n_total == 0:
        return None
    if n_positive == 0:
        if mean_t_air_c is not None:
            floor_per_1000 = dose_response_ehs_probability_floor(mean_t_air_c)
            floor_desc = (
                f"Falmouth's own real-world rate at today's actual "
                f"race-window temperature ({mean_t_air_c:.1f}\u00b0C): "
                f"\u2248{floor_per_1000:.2f} per 1000"
            )
        else:
            floor_per_1000 = 1000.0 * dose_response_ehs_probability(0.0)
            floor_desc = (
                f"the dose-response curve's own floor value "
                f"(\u2248{floor_per_1000:.1f} per 1000)"
            )
        text = (
            f"None of the {n_total} simulated participants ever reached a "
            f"non-zero dose (T_rect\u226540.5\u00b0C AND CO_reserve\u22640, "
            f"simultaneously, at any point). The figure above is not 0 "
            f"because it is {floor_desc}. That floor is not arbitrary: it "
            f"is Falmouth's own published temperature-only regression "
            f"(DeMartini et al. 2014) evaluated at this scenario's own "
            f"temperature, not something this simulation detected or "
            f"measured. It may not transfer well to an event that differs "
            f"from Falmouth in distance, field composition, or country. "
            f"Read it as \u201cno elevated risk found beyond that "
            f"background rate\u201d, not as a literal headcount for this "
            f"event."
        )
        if broad_screen_pct is not None:
            text += (
                f" The broader cross-check (\u201cWorth monitoring, broad "
                f"screen\u201d: T_rect\u226540.5\u00b0C OR dehydration\u22652% "
                f"OR RPE\u226517, NOT gated by this same floor) reads "
                f"{broad_screen_pct:.1f}% for this run \u2014 "
            )
            text += ("also near zero, consistent with nothing detected here."
                     if broad_screen_pct < 5 else
                     "not negligible, worth a closer look even though the "
                     "dose-response figure above is only the floor value.")
        return text
    if n_positive < DOSE_POSITIVE_WARNING_THRESHOLD:
        return (
            f"Based on only {n_positive} of {n_total} simulated participants "
            f"who ever reached a non-zero dose. At this small a count, the "
            f"headline estimate can be dominated by one or two individuals "
            f"rather than reflecting a broad pattern across the group -- it "
            f"may say more about a couple of relatively unfit or hard-pushed "
            f"individuals than about typical risk on this day. Increasing "
            f"'Number of simulations' will show whether this number is a "
            f"stable, repeatable feature of this population or mostly "
            f"small-sample noise."
        )
    return None


def dose_response_ehs_probability_floor(mean_t_air_c: float) -> float:
    """[2026-08-13 patch] What a zero-dose participant is assigned, per
    1000 -- now Falmouth's own real-world temperature-only rate at this
    scenario's own temperature, not a fixed global constant. Imported
    lazily (not at module load) to avoid a hard dependency from
    report_generator.py on hestia_bridge.py for callers that never need
    HESTIA at all (e.g. the pyrox-only test suite)."""
    from hestia_bridge import falmouth_ehs_per_1000
    return falmouth_ehs_per_1000(mean_t_air_c)


def _set_cell_shading(cell, hex_colour: str):
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_colour,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _disable_first_row_style(table):
    """Table styles like 'Light List Accent 1' auto-apply special (often
    white-on-colour) formatting to the first row via w:tblLook/@firstRow.
    That clashes with our own per-cell shading -- disable it so plain
    cell-level shading is the only thing affecting appearance."""
    look = table._tbl.tblPr.find(qn("w:tblLook"))
    if look is not None:
        look.set(qn("w:firstRow"), "0")


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h


def _add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    return p


def _fig_to_png_bytes(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _weather_chart(weather_df, exp_start, finish, post_finish_end=None):
    """T_air/WBGT/UTCI/MRT over the race window, with the same colour
    mapping and flag-zone shading as the live app's Plotly chart, so a
    reader who has seen both recognises them as the same data.

    post_finish_end : optional Timestamp -- if given, the span between
    `finish` and `post_finish_end` is shaded and labelled separately, to
    mark HESTIA's own post-finish simulation window (metabolic afterglow +
    acute venous pooling; PF_DUUR_MIN in hestia_model.py) as distinct from
    the race itself. Purely visual -- does not change which rows are
    plotted (the window margin below already covers it).

    Returns (png_bytes, plotted_labels) -- plotted_labels lists only the
    variables that actually had usable (non-all-NaN) data, so the caption
    can never claim a line is shown when the underlying column was
    missing or empty.
    """
    import matplotlib.dates as mdates

    window_end = (post_finish_end if post_finish_end is not None else finish)
    window = weather_df[
        (weather_df.index >= pd.Timestamp(exp_start) - pd.Timedelta(hours=6))
        & (weather_df.index <= window_end + pd.Timedelta(hours=2))
    ]
    if window.empty or len(window) < 2:
        return None, []

    fig, ax = plt.subplots(figsize=(6.8, 3.2))
    x = window.index.tz_localize(None) if window.index.tz is not None else window.index
    plotted = []
    all_nan_vars = []

    for col, style, colour, lw, label in [
        ("MRT", ":", "#1f77b4", 1.4, "MRT"),
        ("T_air_urban", "-", "#ff7f0e", 1.6, "T_air"),
        ("UTCI", "-", "#9467bd", 1.6, "UTCI"),
        ("WBGT", "-", "#d62728", 1.8, "WBGT"),
    ]:
        if col not in window.columns:
            continue
        y = window[col].to_numpy(dtype=float)
        if np.all(np.isnan(y)):
            # Present as a column but nothing usable in it -- do NOT feed
            # an all-NaN series to ax.plot(): a line with no finite y
            # values anywhere contributes no data to matplotlib's x-axis
            # autoscale either (even though x itself is perfectly valid
            # real timestamps), which is what caused the axis to fall
            # back to a nonsensical multi-year default range instead of
            # the real ~20-hour window. Skipping it here, and anchoring
            # xlim explicitly below regardless, are both needed.
            all_nan_vars.append(label)
            continue
        ax.plot(x, y, style, color=colour, linewidth=lw, label=label)
        plotted.append(label)

    # Always anchor the x-axis to the real timestamp range, regardless of
    # whether any y-data was plottable -- this is what actually prevents
    # the multi-year fallback range, not just skipping all-NaN lines.
    ax.set_xlim(x[0], x[-1])

    if not plotted and not all_nan_vars:
        plt.close(fig)
        return None, []

    ax.axhspan(23, 28, color="#d62728", alpha=0.10, zorder=0)
    ax.axhspan(28, max(45, ax.get_ylim()[1]), color="#450a0a", alpha=0.22, zorder=0)

    start_naive = (pd.Timestamp(exp_start).tz_localize(None)
                   if pd.Timestamp(exp_start).tz is not None else pd.Timestamp(exp_start))
    finish_naive = (pd.Timestamp(finish).tz_localize(None)
                    if pd.Timestamp(finish).tz is not None else pd.Timestamp(finish))
    y_min, y_top = ax.get_ylim()
    label_y_finish = y_min + 0.90 * (y_top - y_min)
    ax.axvline(start_naive, color="#1e293b", linestyle="--", linewidth=1.2)
    ax.text(start_naive, y_top, f" start {start_naive.strftime('%H:%M')}", fontsize=7,
           va="top", ha="left", color="#1e293b", fontweight="bold")
    ax.axvline(finish_naive, color="#1e293b", linestyle=":", linewidth=1.2)
    ax.text(finish_naive, label_y_finish, f" finish {finish_naive.strftime('%H:%M')} ", fontsize=7,
           va="top", ha="left", color="#1e293b", fontweight="bold",
           bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.75))

    if post_finish_end is not None:
        pf_naive = (pd.Timestamp(post_finish_end).tz_localize(None)
                   if pd.Timestamp(post_finish_end).tz is not None
                   else pd.Timestamp(post_finish_end))
        ax.axvspan(finish_naive, pf_naive, color="#0ea5e9", alpha=0.10, zorder=0)
        ax.text(pf_naive, y_min, " post-finish window ", fontsize=6.5,
               va="bottom", ha="right", color="#0c4a6e", style="italic")

    if all_nan_vars:
        ax.text(0.5, 0.06, "No data: " + ", ".join(all_nan_vars),
               transform=ax.transAxes, fontsize=7.5, ha="center", va="bottom",
               color="#7f1d1d",
               bbox=dict(boxstyle="round,pad=0.3", fc="#fee2e2", ec="#7f1d1d", alpha=0.85))

    # Clear HH:MM tick labels, one per hour -- the earlier bare "06", "07"
    # style (date-only ticks with the hour folded in) read ambiguously.
    ax.xaxis.set_major_locator(mdates.HourLocator(interval=max(1, len(x) // 12)))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M"))
    # A single date label under the axis, since every tick now spells out
    # the time only -- avoids repeating the date on every tick.
    ax.set_xlabel(start_naive.strftime("%A %d %B %Y"), fontsize=8)

    ax.set_ylabel("\u00b0C", fontsize=9)
    if plotted:
        ax.legend(loc="upper left", fontsize=7, ncol=4, frameon=False)
    ax.tick_params(labelsize=7)
    fig.autofmt_xdate()
    fig.tight_layout()
    return _fig_to_png_bytes(fig), plotted


def _day_overview_chart(weather_df, exp_date, exp_start=None):
    """T_air/WBGT/UTCI/MRT across the WHOLE chosen calendar day (00:00-
    24:00 local), independent of any one level's pace or start time --
    the day-level picture a reader should see before zooming into a
    specific run. Same visual language as _weather_chart (colours, flag
    shading) so the two read as one report, not two different styles.

    exp_start, if given, is marked with a vertical line so the reader can
    place the chosen start time in the context of the full day.

    Returns (png_bytes, plotted_labels), same contract as _weather_chart.
    """
    import matplotlib.dates as mdates

    day = pd.Timestamp(exp_date)
    tz = weather_df.index.tz
    day_start = pd.Timestamp(day.date()).tz_localize(tz) if tz is not None else pd.Timestamp(day.date())
    day_end = day_start + pd.Timedelta(days=1)
    window = weather_df[(weather_df.index >= day_start) & (weather_df.index < day_end)]
    if window.empty or len(window) < 2:
        return None, []

    fig, ax = plt.subplots(figsize=(6.8, 3.0))
    x = window.index.tz_localize(None) if window.index.tz is not None else window.index
    plotted, all_nan_vars = [], []

    for col, style, colour, lw, label in [
        ("MRT", ":", "#1f77b4", 1.3, "MRT"),
        ("T_air_urban", "-", "#ff7f0e", 1.5, "T_air"),
        ("UTCI", "-", "#9467bd", 1.5, "UTCI"),
        ("WBGT", "-", "#d62728", 1.7, "WBGT"),
    ]:
        if col not in window.columns:
            continue
        y = window[col].to_numpy(dtype=float)
        if np.all(np.isnan(y)):
            all_nan_vars.append(label)
            continue
        ax.plot(x, y, style, color=colour, linewidth=lw, label=label)
        plotted.append(label)

    ax.set_xlim(x[0], x[-1])
    if not plotted and not all_nan_vars:
        plt.close(fig)
        return None, []

    ax.axhspan(23, 28, color="#d62728", alpha=0.10, zorder=0)
    ax.axhspan(28, max(45, ax.get_ylim()[1]), color="#450a0a", alpha=0.22, zorder=0)

    if exp_start is not None:
        start_naive = (pd.Timestamp(exp_start).tz_localize(None)
                       if pd.Timestamp(exp_start).tz is not None else pd.Timestamp(exp_start))
        ax.axvline(start_naive, color="#1e293b", linestyle="--", linewidth=1.1)
        ax.text(start_naive, ax.get_ylim()[1], f" chosen start {start_naive.strftime('%H:%M')}",
               fontsize=7, va="top", ha="left", color="#1e293b", fontweight="bold")

    if all_nan_vars:
        ax.text(0.5, 0.06, "No data: " + ", ".join(all_nan_vars),
               transform=ax.transAxes, fontsize=7.5, ha="center", va="bottom",
               color="#7f1d1d",
               bbox=dict(boxstyle="round,pad=0.3", fc="#fee2e2", ec="#7f1d1d", alpha=0.85))

    ax.xaxis.set_major_locator(mdates.HourLocator(interval=2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M"))
    ax.set_xlabel(day_start.strftime("%A %d %B %Y"), fontsize=8)
    ax.set_ylabel("\u00b0C", fontsize=9)
    if plotted:
        ax.legend(loc="upper left", fontsize=7, ncol=4, frameon=False)
    ax.tick_params(labelsize=7)
    fig.autofmt_xdate()
    fig.tight_layout()
    return _fig_to_png_bytes(fig), plotted


def _pyrox_multiday_chart(pyrox_result, label_for, exp_date):
    """Reserve across the whole computed window, one line per level, with
    a marker on the chosen race date -- gives visual context for the
    single reserve percentage reported in the per-level table."""
    from loop_view import reserve_series

    dates = [pd.Timestamp(d) for d in pyrox_result["dates"]]
    if len(dates) < 2:
        return None

    fig, ax = plt.subplots(figsize=(6.4, 2.6))
    colours = plt.cm.tab10.colors
    any_data_plotted = False
    all_nan_labels = []
    for i, (name, res) in enumerate(pyrox_result["groups"].items()):
        y = reserve_series(res)
        n = min(len(dates), len(y))
        y_slice = np.asarray(y[:n], dtype=float)
        label = label_for.get(name, name)
        if n == 0 or np.all(np.isnan(y_slice)):
            # No plottable data for this group at all -- skip the line
            # (plotting an all-NaN series contributes nothing to
            # matplotlib's axis auto-scaling and, if EVERY group is like
            # this, leaves it with no finite data anywhere to anchor to,
            # which is what caused the axis to fall back to a
            # nonsensical multi-year default range instead of the real
            # dates). Track it to show a visible "no data" note instead
            # of a silently blank chart.
            all_nan_labels.append(label)
            continue
        any_data_plotted = True
        ax.plot(dates[:n], y_slice, "-o", markersize=2.5, linewidth=1.3,
               color=colours[i % len(colours)], label=label)

    # Always anchor the x-axis to the real date range, regardless of
    # whether any y-data was plottable -- this is what actually fixes
    # the multi-year fallback range, not just skipping NaN lines above.
    ax.set_xlim(dates[0], dates[-1])

    ax.axhspan(0, 25, color="#d62728", alpha=0.08, zorder=0)
    ax.axhspan(25, 50, color="#eab308", alpha=0.08, zorder=0)
    target = pd.Timestamp(exp_date)
    ax.axvline(target, color="#1e293b", linestyle="--", linewidth=1)
    ax.text(target, 102, "chosen date", fontsize=7, va="bottom",
           ha="center", color="#1e293b")
    ax.set_ylim(0, 108)
    ax.set_ylabel("% reserve remaining", fontsize=9)
    if any_data_plotted:
        ax.legend(loc="lower left", fontsize=6.5, ncol=2, frameon=False)
    if all_nan_labels:
        note = ("No data: " + ", ".join(all_nan_labels) if any_data_plotted
                else "No data available for this period")
        ax.text(0.5, 0.5, note, transform=ax.transAxes, fontsize=8,
               ha="center", va="center", color="#7f1d1d",
               bbox=dict(boxstyle="round,pad=0.3", fc="#fee2e2", ec="#7f1d1d", alpha=0.85))
    ax.tick_params(labelsize=7)
    fig.autofmt_xdate()
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _hestia_distribution_chart(peak_t_rect_all: list, level_label: str):
    """Histogram of peak T_re across the simulated population -- shows
    the SPREAD behind the summary percentages, which is more honest than
    a single mean/percentage given the PROVISIONAL calibration status."""
    vals = [v for v in peak_t_rect_all if v is not None and not np.isnan(v)]
    if len(vals) < 5:
        return None

    fig, ax = plt.subplots(figsize=(5.0, 2.2))
    ax.hist(vals, bins=min(20, max(5, len(vals) // 3)), color="#9467bd", alpha=0.75)
    ax.axvline(40.5, color="#7f0000", linestyle="--", linewidth=1.2)
    ax.text(40.5, ax.get_ylim()[1], " 40.5\u00b0C", fontsize=7, va="top", color="#7f0000")
    ax.set_xlabel("Peak T_re per simulated participant (\u00b0C)", fontsize=8)
    ax.set_ylabel("count", fontsize=8)
    ax.tick_params(labelsize=7)
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _co_reserve_distribution_chart(worst_co_reserve_all: list, level_label: str):
    """Histogram of each participant's worst CO_reserve (during the race
    or the 10-min post-finish window) -- the spread behind the reported
    'reached zero/negative capacity' percentage."""
    vals = [v for v in worst_co_reserve_all if v is not None and not np.isnan(v)]
    if len(vals) < 5:
        return None

    fig, ax = plt.subplots(figsize=(5.0, 2.2))
    ax.hist(vals, bins=min(20, max(5, len(vals) // 3)), color="#2a9d8f", alpha=0.75)
    ax.axvline(0, color="#7f0000", linestyle="--", linewidth=1.2)
    ax.text(0, ax.get_ylim()[1], " 0 L/min", fontsize=7, va="top", color="#7f0000")
    ax.set_xlabel("Worst CO_reserve per simulated participant (L/min)", fontsize=8)
    ax.set_ylabel("count", fontsize=8)
    ax.tick_params(labelsize=7)
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def dose_evolution_chart(representative_traces: list, level_label: str = ""):
    """T_rect, CO_reserve, and cumulative dose over time. First trace
    (from hestia_bridge's _select_representative_traces) is always the
    real population median, computed point-by-point across every
    simulated participant -- shown distinctly (thick black dashed line)
    from the up-to-3 illustrative individual examples that follow it.
    A thin, colour-matched dotted vertical line marks each trace's own
    stop/finish time -- these genuinely differ between participants
    (faster/slower pacing), which is also where the sharp post-finish
    CO_reserve recovery in the CO_reserve panel comes from. Shared
    between the Word report and any live Streamlit page."""
    if not representative_traces:
        return None

    individual_colours = ["#2a9d8f", "#e9a942", "#c1121f"]
    fig, axes = plt.subplots(3, 1, figsize=(7.2, 8.0), sharex=True)

    individual_i = 0
    stop_markers = []  # (colour, stopped_at) pairs, drawn after the main lines
    for tr in representative_traces:
        if tr["label"] == "Population median":
            style = dict(color="#1e293b", linestyle="--", linewidth=2.2, zorder=5)
        else:
            style = dict(color=individual_colours[individual_i % len(individual_colours)],
                        linewidth=1.4, alpha=0.85)
            individual_i += 1
        axes[0].plot(tr["min"], tr["t"], label=tr["label"], **style)
        axes[1].plot(tr["min"], tr["c"], **style)
        axes[2].plot(tr["min"], tr["dose"], **style)
        if tr.get("stopped_at"):
            stop_markers.append((style["color"], tr["stopped_at"]))

    for colour, stopped_at in stop_markers:
        for ax in axes:
            ax.axvline(stopped_at, color=colour, linestyle=":", linewidth=1.1, alpha=0.55)

    axes[0].axhline(40.5, color="#7f1d1d", linestyle="--", linewidth=1, alpha=0.7)
    axes[0].text(0, 40.5, " 40.5\u00b0C", fontsize=7, va="bottom", color="#7f1d1d")
    axes[0].set_ylabel("T_rect (\u00b0C)")
    axes[0].legend(loc="upper left", fontsize=8, frameon=False)
    title = "How risk builds over time" + (f" \u2014 {level_label}" if level_label else "")
    axes[0].set_title(title)

    axes[1].axhline(0, color="#7f1d1d", linestyle="--", linewidth=1, alpha=0.7)
    axes[1].set_ylabel("CO_reserve (L/min)")

    axes[2].set_ylabel("Cumulative dose Q\n(L/min-minutes)")
    axes[2].set_xlabel("Minutes since start (race + post-finish)")
    axes[2].text(0.5, -0.32, "Dotted vertical lines: each line's own stop/finish time",
                fontsize=7, color="#475569", ha="center", transform=axes[2].transAxes)

    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _t_rect_co_reserve_scatter(pairs: list, level_label: str):
    """T_rect vs CO_reserve, one point per (participant, timestep) across
    the race and the post-finish window -- the same underlying data the
    'true EHS criterion' percentage is computed from, so every point that
    falls in the shaded quadrant is literally a moment that criterion
    counted. Not an approximation from two separately-extracted maxima."""
    pairs = [(t, c) for t, c in pairs if t is not None and c is not None
            and not np.isnan(t) and not np.isnan(c)]
    if len(pairs) < 10:
        return None

    t_vals = np.array([p[0] for p in pairs])
    c_vals = np.array([p[1] for p in pairs])

    fig, ax = plt.subplots(figsize=(5.6, 4.2))

    x_min, x_max = min(37.0, t_vals.min() - 0.2), max(41.5, t_vals.max() + 0.2)
    y_min, y_max = min(-1.0, c_vals.min() - 0.3), max(3.0, c_vals.max() + 0.3)

    # The danger quadrant: T_rect >= 40.5 AND CO_reserve <= 0 -- exactly
    # the true EHS criterion this app uses.
    ax.axvspan(40.5, x_max, ymin=0, ymax=(0 - y_min) / (y_max - y_min),
              color="#7f0000", alpha=0.18, zorder=0)
    ax.axhline(0, color="#94a3b8", linewidth=0.8, zorder=1)
    ax.axvline(40.5, color="#94a3b8", linewidth=0.8, zorder=1)

    in_quadrant = (t_vals >= 40.5) & (c_vals <= 0)
    ax.scatter(t_vals[~in_quadrant], c_vals[~in_quadrant], s=7, alpha=0.35,
              color="#1f77b4", linewidths=0, label="Outside criterion")
    ax.scatter(t_vals[in_quadrant], c_vals[in_quadrant], s=10, alpha=0.75,
              color="#7f0000", linewidths=0, label="Meets true EHS criterion")

    ax.text(x_max - 0.05, y_min + 0.1,
           f"n={int(in_quadrant.sum())} of {len(pairs)} timestep-points\nin this quadrant "
           "(not the\nsame count as participants \u2014 see caption)",
           fontsize=6.5, ha="right", va="bottom", color="#7f0000")

    ax.set_xlim(x_min, x_max)
    ax.set_ylim(y_min, y_max)
    ax.set_xlabel("T_rect (\u00b0C)", fontsize=9)
    ax.set_ylabel("CO_reserve (L/min)", fontsize=9)
    ax.legend(loc="upper left", fontsize=7, frameon=False)
    ax.tick_params(labelsize=7)
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _flag_colour(flag_name: str) -> str:
    return {
        "Green (low)": "C6EFCE", "Yellow (moderate)": "FFEB9C",
        "Red (high)": "FFC7CE", "Black (extreme)": "8B0000",
        "no data": "F2F2F2",
    }.get(flag_name, "F2F2F2")


# =============================================================================
# Section builders
# =============================================================================
def _add_title_section(doc, city_name, exp_start, generated_at, app_build,
                       report_title="PYROX \u2014 Findings Report"):
    title = doc.add_heading(report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run(f"{city_name} \u2014 {pd.Timestamp(exp_start).strftime('%A %d %B %Y, %H:%M')}").bold = True

    meta = doc.add_paragraph()
    meta.add_run(
        f"Report generated: {generated_at.strftime('%Y-%m-%d %H:%M %Z')}\n"
        f"App build: {app_build}"
    ).font.size = Pt(9)

    scope = doc.add_paragraph()
    scope_run = scope.add_run(
        "SCOPE: This report presents model outputs and findings only. It "
        "does not include recommendations for operational measures "
        "(staffing, start times, water stations, or any other action) \u2014 "
        "those decisions rest with the organising and medical team, who "
        "have context this report does not. Where two figures in this "
        "report disagree, an explanation of WHY is included where "
        "relevant; that is a methodological note, not a recommendation."
    )
    scope_run.italic = True
    scope_run.font.size = Pt(9.5)
    scope.paragraph_format.space_before = Pt(6)
    scope.paragraph_format.space_after = Pt(12)

    doc.add_paragraph().add_run("").add_break()


def _add_weather_section(doc, weather_df, exp_start, finish):
    _add_heading(doc, "Weather conditions used", level=1)
    window = weather_df[
        (weather_df.index >= pd.Timestamp(exp_start) - pd.Timedelta(hours=2))
        & (weather_df.index <= finish + pd.Timedelta(hours=1))
    ]
    if window.empty:
        doc.add_paragraph("No weather data available for this window.")
        return

    cols = [c for c in ["T_air_urban", "WBGT", "UTCI", "MRT"] if c in window.columns]
    table = doc.add_table(rows=1, cols=1 + len(cols))
    table.style = "Light Grid Accent 1"
    _disable_first_row_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    labels = {"T_air_urban": "T_air (\u00b0C)", "WBGT": "WBGT (\u00b0C)",
             "UTCI": "UTCI (\u00b0C)", "MRT": "MRT (\u00b0C)"}
    for i, c in enumerate(cols):
        hdr[i + 1].text = labels.get(c, c)

    for stat_name, fn in [("Peak", "max"), ("Mean over window", "mean"), ("Minimum", "min")]:
        row = table.add_row().cells
        row[0].text = stat_name
        for i, c in enumerate(cols):
            val = getattr(window[c], fn)()
            row[i + 1].text = f"{val:.1f}"

    _add_caption(doc, "Window: 2h before the chosen start time to 1h after the "
                      "estimated finish time, from the weather data used for this run.")
    doc.add_paragraph()

    chart, plotted = _weather_chart(weather_df, exp_start, finish)
    if chart is not None:
        doc.add_picture(chart, width=Cm(15.5))
        var_list = ", ".join(plotted[:-1]) + (" and " + plotted[-1] if len(plotted) > 1 else plotted[0])
        _add_caption(doc, f"{var_list}, hour by hour, from 2h before the chosen "
                          "start to 1h after the estimated finish -- shown wider "
                          "than the race itself so the trend leading in and out "
                          "is visible. Dashed line = start, dotted line = "
                          "estimated finish. Shaded bands mark the red "
                          "(23\u201328\u00b0C) and black (>28\u00b0C) WBGT flag zones.")
        missing = [v for v in ("WBGT", "UTCI") if v not in plotted]
        if missing:
            _add_caption(doc, f"\u26a0\ufe0f {' and '.join(missing)} data was not "
                              "available for this run and is not shown above.")
        doc.add_paragraph()


def _add_per_level_section(doc, per_level_exposure, pyrox_result, label_for,
                           hestia_results, exp_date, exp_start=None, finish=None,
                           met_by_level_label=None, duration_by_level_label=None):
    from loop_view import reserve_series
    met_by_level_label = met_by_level_label or {}
    duration_by_level_label = duration_by_level_label or {}
    # Fallback only: the shared report-wide finish, used solely when no
    # per-level duration was supplied (e.g. older callers). Using this for
    # every level is what caused the original bug -- each level normally
    # has its own duration, which should always be preferred when present.
    fallback_minutes = ((finish - exp_start).total_seconds() / 60.0
                        if exp_start is not None and finish is not None else None)
    from decision_support import worst_flag, flag_display_name

    _add_heading(doc, "Per-level findings", level=1)

    multiday_chart = _pyrox_multiday_chart(pyrox_result, label_for, exp_date)
    if multiday_chart is not None:
        doc.add_picture(multiday_chart, width=Cm(15.5))
        _add_caption(doc, "PYROX multi-day reserve across the whole computed window, "
                          "for context: where the chosen date sits relative to the "
                          "days before and after it.")
        doc.add_paragraph()

    for name, res in pyrox_result["groups"].items():
        level_label = label_for.get(name, name)
        exposure = per_level_exposure.get(level_label)
        flag = worst_flag(exposure) if exposure else None
        flag_name = flag_display_name(flag) if flag else "no data"

        dates = pyrox_result["dates"]
        target = pd.Timestamp(exp_date).date()
        matches = [i for i, d in enumerate(dates) if pd.Timestamp(d).date() == target]
        reserve_pct = reserve_series(res)[matches[0]] if matches else float("nan")

        _add_heading(doc, level_label, level=2)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        _disable_first_row_style(table)

        def _row(label, value, shade=None):
            r = table.add_row().cells
            r[0].text = label
            r[1].text = str(value)
            if shade:
                _set_cell_shading(r[1], shade)
                for para in r[1].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                        run.font.bold = True

        _row("Worst WBGT flag reached (race window)", flag_name, _flag_colour(flag_name))
        _row("PYROX multi-day reserve, on this date",
             f"{reserve_pct:.0f}%" if not pd.isna(reserve_pct) else "no data")

        hestia = hestia_results.get(level_label)
        if hestia is not None:
            pinned = hestia.get("pct_vo2max_pinned")
            falmouth_est = hestia.get("falmouth_ehs_per_1000")
            mean_t = hestia.get("mean_t_air_race_window")
            dose_pct = hestia.get("pct_dose_response_ehs")
            level_met = met_by_level_label.get(level_label)

            ehs_ci = hestia.get("ehs_interval")
            if dose_pct is not None:
                _ehs_val = f"\u2248{dose_pct*10:.1f} per 1000"
                if ehs_ci and "error" not in ehs_ci:
                    _pct = int(round((1 - ehs_ci["alpha"]) * 100))
                    _ehs_val += (
                        f"  \u2014  {_pct}% sampling + anchor interval "
                        f"{ehs_ci['lo_per_1000']:.2f}\u2013"
                        f"{ehs_ci['hi_per_1000']:.2f} per 1000"
                    )
                _row("EHS estimate (primary: dose-response model, see note below)",
                     _ehs_val)
                race_minutes = duration_by_level_label.get(level_label, fallback_minutes)
                met_off = level_met is not None and abs(level_met - 10.5) > 3.0
                dur_off = race_minutes is not None and abs(race_minutes - 96) > 60
                if met_off or dur_off:
                    warn = doc.add_paragraph()
                    warn_run = warn.add_run(
                        "\u26a0\ufe0f This level "
                        + (f"(MET {level_met:.1f}, " if level_met is not None else "(")
                        + (f"{race_minutes:.0f} min) " if race_minutes is not None else ") ")
                        + "falls outside the range the dose-response curve was "
                          "actually fit on (MET\u224810.5, \u224896 min). Not validated "
                          "to generalise this far -- treat with extra caution."
                    )
                    warn_run.italic = True
                    warn_run.font.color.rgb = RGBColor(0x7f, 0x1d, 0x1d)

            _row("HESTIA \u2014 peak T_re, mean", f"{hestia['peak_t_rect_mean']:.1f}\u00b0C")
            _row("HESTIA (raw, uncalibrated) \u2014 true EHS criterion met "
                 "(T_rect\u226540.5\u00b0C AND CO_reserve\u22640, simultaneous)",
                 f"{hestia['pct_true_ehs_criterion']:.1f}% "
                 f"(\u2248{round(hestia['pct_true_ehs_criterion'] * 10):.0f} per 1000, "
                 "NOT the calibrated estimate above)")
            _row("HESTIA \u2014 broad monitoring screen (not a medical incident rate)",
                 f"{hestia['pct_first_aid']:.1f}%")
            _row("HESTIA \u2014 avg. cardiovascular capacity remaining",
                 f"{hestia['pct_reserve_remaining_mean']:.0f}%")
            _row("HESTIA (raw, uncalibrated) \u2014 reached zero/negative capacity",
                 f"{hestia['pct_zero_or_negative_capacity']:.1f}%")
            if pinned is not None and not pd.isna(pinned) and pinned > 20:
                _row("HESTIA \u2014 VO2max-pinning caution",
                     f"{pinned:.0f}% of the simulated population at effort ceiling")

            if dose_pct is not None:
                summary = doc.add_paragraph()
                summary.add_run(
                    f"Per 1000 participants at this level, under these "
                    f"conditions: \u2248{dose_pct*10:.1f} are estimated to "
                    f"experience exertional heat stroke, from a "
                    f"dose-response model over this scenario's actual "
                    f"simulated pace, duration and group."
                ).italic = True
                n_dose_pos, n_dose_total = dose_positive_count(hestia)
                note_parts = []
                if falmouth_est is not None:
                    note_parts.append(
                        f"epidemiologically-calibrated estimate (Falmouth, "
                        f"temperature-only, DeMartini et al. 2014): "
                        f"\u2248{falmouth_est:.1f} per 1000"
                        + (f" at {mean_t:.1f}\u00b0C" if mean_t is not None else "")
                    )
                note_parts.append(
                    f"raw HESTIA simulation (uncalibrated): "
                    f"{hestia['pct_true_ehs_criterion']:.1f}% "
                    f"(\u2248{round(hestia['pct_true_ehs_criterion']*10):.0f} per 1000)"
                )
                if n_dose_total > 0:
                    note_parts.append(
                        f"based on {n_dose_pos} of {n_dose_total} simulated "
                        f"participants (distinct participants, not "
                        f"timestep-points) who ever reached a non-zero dose"
                    )
                _add_caption(doc, "For comparison \u2014 " + "; ".join(note_parts) + ". "
                                  "The dose-response model is a logistic curve over "
                                  "each participant's cumulative T_rect/CO_reserve "
                                  "deficit (depth \u00d7 duration), fit jointly against "
                                  "the Falmouth data across 5 temperature scenarios, "
                                  "refit 2026-08-10 after a clo_value correction that "
                                  "fixed a major T_rect over-prediction (predicted/"
                                  "target ratio now 0.6-1.8x, wider than pre-fix, "
                                  "because far fewer simulated participants now enter "
                                  "the danger quadrant, leaving less data to fit). "
                                  "EXPLORATORY: fit at n=120/scenario, well below "
                                  "production scale.")
                if ehs_ci and "error" not in ehs_ci:
                    from uncertainty import interval_caveats
                    for _c in interval_caveats(ehs_ci):
                        _p = doc.add_paragraph()
                        _r = _p.add_run("\u26a0\ufe0f " + _c)
                        _r.italic = True
                        _r.font.color.rgb = RGBColor(0x7f, 0x1d, 0x1d)

                dose_warning = dose_positive_warning_text(
                    n_dose_pos, n_dose_total, hestia.get("pct_first_aid") if hestia else None,
                    mean_t)
                if dose_warning:
                    icon = "\u2139\ufe0f" if n_dose_pos == 0 else "\u26a0\ufe0f"
                    warn_p = doc.add_paragraph()
                    warn_run = warn_p.add_run(icon + " " + dose_warning)
                    warn_run.italic = True
                    warn_run.font.color.rgb = (RGBColor(0x1e, 0x40, 0xaf) if n_dose_pos == 0
                                               else RGBColor(0x7f, 0x1d, 0x1d))

            dist_chart = _hestia_distribution_chart(
                hestia.get("peak_t_rect_all", []), level_label)
            if dist_chart is not None:
                doc.add_paragraph()
                doc.add_picture(dist_chart, width=Cm(11.5))
                _add_caption(doc, "Distribution of peak T_re across the simulated "
                                  "population for this level -- the spread behind the "
                                  "summary figures above. Dashed line marks the "
                                  "40.5\u00b0C reference used in the true EHS criterion.")

            co_chart = _co_reserve_distribution_chart(
                hestia.get("worst_co_reserve_all", []), level_label)
            if co_chart is not None:
                doc.add_paragraph()
                doc.add_picture(co_chart, width=Cm(11.5))
                _add_caption(doc, "Distribution of each participant's worst "
                                  "cardiovascular reserve (CO_reserve) during the race "
                                  "or the post-finish window -- the spread behind the "
                                  "'reached zero/negative capacity' figure above.")

            scatter = _t_rect_co_reserve_scatter(
                hestia.get("t_rect_co_reserve_pairs", []), level_label)
            if scatter is not None:
                doc.add_paragraph()
                doc.add_picture(scatter, width=Cm(13.5))
                _add_caption(doc, "T_rect against CO_reserve, every simulated "
                                  "participant at every timestep (race and "
                                  "post-finish window). The shaded quadrant "
                                  "(T_rect\u226540.5\u00b0C AND CO_reserve\u22640) is exactly "
                                  "the true EHS criterion reported in the table above -- "
                                  "this plot shows the same data point by point, not an "
                                  "approximation from separate maxima. Note: the point "
                                  "count in the quadrant is per TIMESTEP, not per "
                                  "participant -- one participant who stays in the "
                                  "quadrant for several consecutive timesteps "
                                  "contributes several points, so this count is not "
                                  "directly comparable to the participant-level "
                                  "percentage above.")
        else:
            _row("HESTIA acute capacity data", "not calculated for this level in this run")

        doc.add_paragraph()


def _add_divergence_section(doc, flag_warnings, hestia_warnings):
    if not flag_warnings and not hestia_warnings:
        return
    _add_heading(doc, "Where figures in this report disagree, and why", level=1)
    doc.add_paragraph(
        "The following is a methodological explanation of why certain "
        "figures differ \u2014 both are correct; they answer different "
        "questions at different timescales. This section describes the "
        "disagreement, not what to do about it."
    )
    for w in flag_warnings:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(w)
    for w in hestia_warnings:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(w)
    doc.add_paragraph()


def _add_limitations_section(doc, hestia_results):
    _add_heading(doc, "Limitations", level=1)
    items = [
        "All figures are modelled group averages, not measurements of or "
        "predictions for any individual.",
        "PYROX's population tier has no event-level validation against "
        "real incident records.",
    ]
    if any(hestia_results.values()):
        items.append(
            "The 'EHS estimate (epidemiologically calibrated)' figure uses "
            "a published regression against real Falmouth Road Race "
            "incident data (DeMartini et al. 2014), not HESTIA's own "
            "physiological simulation directly. HESTIA's raw simulation "
            "was found, in testing, to over-predict this same benchmark "
            "by roughly 20-50x; that raw output is still shown (clearly "
            "marked) for transparency, not as the primary figure to use. "
            "The Falmouth regression (R\u00b2=0.65) was fitted on one "
            "specific 7-mile race with a broad recreational-to-elite "
            "field; applying it to a different distance, duration, or "
            "population is itself an approximation, not a validated "
            "transfer."
        )
        items.append(
            "HESTIA's own internal calibration is self-labelled "
            "PROVISIONAL in the model's own source: fit at a reduced "
            "sample size (N=200, not production scale), following a "
            "rebuild of the cardiovascular module. Treat HESTIA's raw "
            "percentages as directional, not as settled probabilities."
        )
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()


# =============================================================================
# Entry point
# =============================================================================
def _add_day_overview_section(doc, weather_df, exp_date, exp_start):
    """The chosen day, in general -- before zooming into any one level's
    start time. Uses the WHOLE calendar day (00:00-24:00 local), not the
    race window, so a reader sees how the day as a whole shaped up before
    reading about the specific run."""
    from decision_support import exposure_by_flag, flag_display_name

    _add_heading(doc, "Day overview", level=1)
    day = pd.Timestamp(exp_date)
    tz = weather_df.index.tz
    day_start = pd.Timestamp(day.date()).tz_localize(tz) if tz is not None else pd.Timestamp(day.date())
    day_end = day_start + pd.Timedelta(days=1)
    window = weather_df[(weather_df.index >= day_start) & (weather_df.index < day_end)]

    if window.empty:
        doc.add_paragraph("No weather data available for this day.")
        return

    cols = [c for c in ["T_air_urban", "WBGT", "UTCI", "MRT"] if c in window.columns]
    table = doc.add_table(rows=1, cols=1 + len(cols))
    table.style = "Light Grid Accent 1"
    _disable_first_row_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    labels = {"T_air_urban": "T_air (\u00b0C)", "WBGT": "WBGT (\u00b0C)",
             "UTCI": "UTCI (\u00b0C)", "MRT": "MRT (\u00b0C)"}
    for i, c in enumerate(cols):
        hdr[i + 1].text = labels.get(c, c)
    for stat_name, fn in [("Peak", "max"), ("Mean over the day", "mean"), ("Minimum", "min")]:
        row = table.add_row().cells
        row[0].text = stat_name
        for i, c in enumerate(cols):
            val = getattr(window[c], fn)()
            row[i + 1].text = f"{val:.1f}"

    _add_caption(doc, f"Full calendar day: {day_start.strftime('%A %d %B %Y')}, "
                      "00:00\u201324:00 local time -- independent of any one "
                      "level's own start time or pace.")
    doc.add_paragraph()

    day_exposure = exposure_by_flag(weather_df, day_start, day_end)
    if day_exposure:
        flag_table = doc.add_table(rows=1, cols=4)
        flag_table.style = "Light List Accent 1"
        _disable_first_row_style(flag_table)
        hdr2 = flag_table.rows[0].cells
        for i, status in enumerate(["race_green", "race_yellow", "race_red", "race_black"]):
            hdr2[i].text = flag_display_name(status)
        row2 = flag_table.add_row().cells
        for i, status in enumerate(["race_green", "race_yellow", "race_red", "race_black"]):
            hrs = day_exposure.get(status, 0.0)
            row2[i].text = f"{hrs:.1f}h"
            if hrs > 0:
                _set_cell_shading(row2[i], _flag_colour(flag_display_name(status)))
        _add_caption(doc, "Hours the whole day spends in each WBGT flag category "
                          "(same thresholds as the athletics-federation flags used "
                          "throughout this suite) -- not specific to any level's "
                          "own exposure, which is shown separately for the chosen "
                          "start time below.")
        doc.add_paragraph()

    chart, plotted = _day_overview_chart(weather_df, exp_date, exp_start)
    if chart is not None:
        doc.add_picture(chart, width=Cm(15.5))
        var_list = ", ".join(plotted[:-1]) + (" and " + plotted[-1] if len(plotted) > 1 else plotted[0])
        _add_caption(doc, f"{var_list}, hour by hour, across the full day. Dashed "
                          "line marks the chosen start time (zoomed in below). "
                          "Shaded bands mark the red (23\u201328\u00b0C) and black "
                          "(>28\u00b0C) WBGT flag zones.")
        doc.add_paragraph()


def _add_race_window_section(doc, level_label, weather_df, exp_start, finish,
                             post_finish_end, hestia_result, met_value=None,
                             duration_minutes=None):
    """Zoomed-in section for one level: the chosen start time through the
    race itself and HESTIA's own post-finish simulation window. Mirrors
    exactly what app_beleid.py shows on screen for this level -- no raw/
    uncalibrated figures, no PROVISIONAL-calibration table rows, no
    multi-day PYROX context. If it isn't on the policy view's screen, it
    isn't in this section either, so the downloaded report never says
    more than the app the reader was just looking at."""
    from decision_support import exposure_by_flag, flag_display_name, worst_flag

    _add_heading(doc, level_label, level=1)

    meta = doc.add_paragraph()
    meta_bits = [f"Start {pd.Timestamp(exp_start).strftime('%H:%M')}",
                f"finish (est.) {pd.Timestamp(finish).strftime('%H:%M')}"]
    if duration_minutes is not None:
        meta_bits.append(f"{duration_minutes:.0f} min")
    if met_value is not None:
        meta_bits.append(f"MET {met_value:.1f}")
    meta.add_run(" \u2014 ".join(meta_bits)).italic = True

    exposure = exposure_by_flag(weather_df, exp_start, finish)
    flag = worst_flag(exposure) if exposure else None
    flag_name = flag_display_name(flag) if flag else "no data"

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    _disable_first_row_style(table)

    def _row(label, value, shade=None):
        r = table.add_row().cells
        r[0].text = label
        r[1].text = str(value)
        if shade:
            _set_cell_shading(r[1], shade)
            for para in r[1].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                    run.font.bold = True

    _row("Worst WBGT flag reached (race window)", flag_name, _flag_colour(flag_name))

    dose_pct = hestia_result.get("pct_dose_response_ehs") if hestia_result else None
    falmouth_est = hestia_result.get("falmouth_ehs_per_1000") if hestia_result else None
    mean_t = hestia_result.get("mean_t_air_race_window") if hestia_result else None
    broad_screen = hestia_result.get("pct_first_aid") if hestia_result else None
    ehs_ci = hestia_result.get("ehs_interval") if hestia_result else None
    if dose_pct is not None:
        _ehs_val = f"\u2248{dose_pct*10:.1f} per 1000"
        if ehs_ci and "error" not in ehs_ci:
            _pct = int(round((1 - ehs_ci["alpha"]) * 100))
            _ehs_val += (
                f"  \u2014  {_pct}% sampling + anchor interval "
                f"{ehs_ci['lo_per_1000']:.2f}\u2013"
                f"{ehs_ci['hi_per_1000']:.2f} per 1000"
            )
        _row("EHS estimate (dose-response model)", _ehs_val)
    if broad_screen is not None:
        _row("Worth monitoring (broad screen: T_rect\u226540.5\u00b0C OR "
             "dehydration\u22652% OR RPE\u226517)", f"{broad_screen:.1f}%")

    doc.add_paragraph()

    chart, plotted = _weather_chart(weather_df, exp_start, finish, post_finish_end)
    if chart is not None:
        doc.add_picture(chart, width=Cm(15.5))
        var_list = ", ".join(plotted[:-1]) + (" and " + plotted[-1] if len(plotted) > 1 else plotted[0])
        _add_caption(doc, f"{var_list}, hour by hour, around this level's own start "
                          "and finish. Dashed line = start, dotted line = estimated "
                          "finish. The lightly shaded band after the finish marks "
                          "HESTIA's post-finish simulation window (metabolic "
                          "afterglow and acute venous pooling continue to raise "
                          "risk for a short time after runners stop -- Roberts "
                          "1998; Rowell 1974).")
        doc.add_paragraph()

    if hestia_result is None:
        doc.add_paragraph(
            "Physiological simulation was not calculated for this level in "
            "this run (the \u201cCalculate\u201d button was not used before "
            "downloading this report)."
        )
        doc.add_paragraph()
        return

    if dose_pct is not None:
        summary = doc.add_paragraph()
        summary.add_run(
            f"Per 1000 participants at this level, under these conditions: "
            f"\u2248{dose_pct*10:.1f} are estimated to experience exertional "
            f"heat stroke, from a dose-response model over this scenario's "
            f"actual simulated pace, duration and group."
        ).italic = True
        n_dose_pos, n_dose_total = dose_positive_count(hestia_result)
        note_parts = []
        if falmouth_est is not None:
            note_parts.append(
                f"alternative estimate, temperature only (Falmouth Road Race, "
                f"DeMartini et al. 2014): \u2248{falmouth_est:.1f} per 1000"
                + (f" at {mean_t:.1f}\u00b0C" if mean_t is not None else "")
            )
        note_text = ("For comparison \u2014 " + "; ".join(note_parts) + ". "
                     if note_parts else "")
        if n_dose_total > 0:
            note_text += (
                f"Based on {n_dose_pos} of {n_dose_total} simulated "
                f"participants (distinct participants, not timestep-points) "
                f"who ever reached a non-zero dose. "
            )
        _add_caption(doc, note_text +
                          "EXPLORATORY calibration: fit at n=120 simulated "
                          "participants per scenario, well below production "
                          "scale. For the full methodology, raw/uncalibrated "
                          "figures, and multi-day PYROX context, see the PYROX "
                          "Participants view.")
        if ehs_ci and "error" not in ehs_ci:
            from uncertainty import interval_caveats
            for _c in interval_caveats(ehs_ci):
                _p = doc.add_paragraph()
                _r = _p.add_run("\u26a0\ufe0f " + _c)
                _r.italic = True
                _r.font.color.rgb = RGBColor(0x7f, 0x1d, 0x1d)
        dose_warning = dose_positive_warning_text(
            n_dose_pos, n_dose_total,
            hestia_result.get("pct_first_aid") if hestia_result else None,
            mean_t)
        if dose_warning:
            icon = "\u2139\ufe0f" if n_dose_pos == 0 else "\u26a0\ufe0f"
            warn_p = doc.add_paragraph()
            warn_run = warn_p.add_run(icon + " " + dose_warning)
            warn_run.italic = True
            warn_run.font.color.rgb = (RGBColor(0x1e, 0x40, 0xaf) if n_dose_pos == 0
                                       else RGBColor(0x7f, 0x1d, 0x1d))

    chart_specs = [
        (_t_rect_co_reserve_scatter(hestia_result.get("t_rect_co_reserve_pairs", []), level_label),
         "T_rect against CO_reserve, every simulated participant at every "
         "timestep (race and post-finish window). The shaded quadrant "
         "(T_rect\u226540.5\u00b0C AND CO_reserve\u22640) is the true EHS criterion "
         "behind the estimate above."),
        (_hestia_distribution_chart(hestia_result.get("peak_t_rect_all", []), level_label),
         "Distribution of peak core temperature across the simulated population."),
        (_co_reserve_distribution_chart(hestia_result.get("worst_co_reserve_all", []), level_label),
         "Distribution of worst cardiovascular reserve across the simulated population."),
    ]
    traces = [tr for tr in hestia_result.get("representative_traces", [])
             if tr.get("label") != "Population median"]
    if traces:
        chart_specs.append((
            dose_evolution_chart(traces, level_label),
            "How risk builds over time for representative participants. "
            "Dotted vertical line: this level's own stop/finish time.",
        ))

    for chart_img, caption in chart_specs:
        if chart_img is not None:
            doc.add_picture(chart_img, width=Cm(13.5))
            _add_caption(doc, caption)
            doc.add_paragraph()


def generate_policy_prediction_excel_bytes(
    city_name: str, lat: float, lon: float, tz_name: str,
    exp_date, exp_start, weather_df: pd.DataFrame, levels_data: dict,
    app_build: str, is_hindcast: bool = False,
) -> bytes:
    """A single-file PRE-EVENT prediction record for app_beleid.py, built
    for LATER comparison against what actually happened -- mirrors
    app_athletes.py's prediction_record_excel_bytes() in purpose, but
    reads from levels_data (this app's own data shape) instead of a
    PYROX multi-day result, for the same reason generate_policy_report_docx
    is a separate function rather than a wrapper: app_beleid.py never
    computes PYROX multi-day results, by design.

    Why this exists: Streamlit Cloud's filesystem is ephemeral -- nothing
    written server-side survives a reboot -- so without something the
    reader downloads themselves, every prediction is lost the moment the
    tab closes. This is also the concrete path from PROVISIONAL
    calibration to something stronger, including for the dose-response
    floor value: the ONLY way to check whether Falmouth's own background
    incidence (baked into the curve when no participant reaches a
    non-zero dose) actually holds for events like this one is to compare
    it against real outcomes, gathered forward across many events rather
    than asserted from one simulation.

    Deliberately includes empty actual_* columns, so the SAME file can be
    completed after the event rather than needing a second document.
    """
    from decision_support import exposure_by_flag, worst_flag, flag_display_name

    buf = io.BytesIO()

    def _strip_tz(df):
        df = df.copy()
        if df.index.tz is not None:
            df.index = df.index.tz_localize(None)
        return df

    made_at = pd.Timestamp.now(tz=tz_name)
    meta = {
        "record_type": "HINDCAST (observed weather, past event)" if is_hindcast
                       else "PRE-EVENT PREDICTION (forecast weather)",
        "prediction_made_at": made_at.strftime("%Y-%m-%d %H:%M %Z"),
        "app_build": app_build,
        "city": city_name, "lat": lat, "lon": lon, "timezone": tz_name,
        "race_date": str(exp_date), "race_start_time": pd.Timestamp(exp_start).strftime("%H:%M"),
        "note": (("This is a HINDCAST run using observed historical weather "
                 "for a date that has already passed -- fill in the "
                 "'actual_*' columns from the real, known outcome to check "
                 "this model's prediction against it directly. "
                 if is_hindcast else
                 "This is a PRE-EVENT prediction from PYROX Beleid. ") +
                "Fill in the 'actual_*' columns in the Per_Level sheet once "
                "real outcomes are known, to build a comparison record over "
                "time -- this is the only way to check whether the "
                "dose-response floor value (Falmouth's own background "
                "incidence, used when no simulated participant reaches a "
                "non-zero dose) actually holds for events like this one."),
    }

    rows = []
    for level_label, d in levels_data.items():
        finish = d["finish"]
        hestia = d.get("hestia_result")
        exposure = exposure_by_flag(weather_df, exp_start, finish)
        flag = worst_flag(exposure) if exposure else None
        n_dose_pos, n_dose_total = dose_positive_count(hestia)

        rows.append({
            "level": level_label,
            "start_time": pd.Timestamp(exp_start).strftime("%H:%M"),
            "finish_time_est": pd.Timestamp(finish).strftime("%H:%M"),
            "duration_minutes": d.get("duration_minutes"),
            "met_value": d.get("met_value"),
            "worst_wbgt_flag": flag_display_name(flag) if flag else "no data",
            "ehs_estimate_dose_response_per_1000":
                round(hestia["pct_dose_response_ehs"] * 10, 2)
                if hestia and hestia.get("pct_dose_response_ehs") is not None else None,
            "alternative_estimate_temp_only_per_1000":
                round(hestia["falmouth_ehs_per_1000"], 2)
                if hestia and hestia.get("falmouth_ehs_per_1000") is not None else None,
            "n_participants_with_nonzero_dose": n_dose_pos if hestia else None,
            "n_simulated_participants": n_dose_total if hestia else None,
            "hestia_broad_screen_pct": round(hestia["pct_first_aid"], 1) if hestia else None,
            "hestia_peak_t_re_mean_c": round(hestia["peak_t_rect_mean"], 2) if hestia else None,
            "hestia_avg_capacity_remaining_pct":
                round(hestia["pct_reserve_remaining_mean"], 1) if hestia else None,
            "hestia_calibration_status":
                "PROVISIONAL (dose-response, n=120/scenario in the fit, not "
                "production-scale)" if hestia else "not calculated for this level in this run",
            # Deliberately empty -- fill in once the real outcome is known.
            "actual_first_aid_visits": None,
            "actual_hospitalisations": None,
            "actual_ehs_cases": None,
            "actual_participant_count": None,
            "notes": None,
        })
    per_level_df = pd.DataFrame(rows)

    weather_window = weather_df[
        (weather_df.index >= pd.Timestamp(exp_start) - pd.Timedelta(hours=6))
        & (weather_df.index <= pd.Timestamp(exp_start) + pd.Timedelta(hours=6))
    ]

    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        pd.DataFrame([meta]).to_excel(writer, sheet_name="Metadata", index=False)
        per_level_df.to_excel(writer, sheet_name="Per_Level", index=False)
        _strip_tz(weather_window).to_excel(writer, sheet_name="Weather_Used")

    return buf.getvalue()


def generate_policy_report_docx(
    city_name: str, exp_date, exp_start, weather_df: pd.DataFrame,
    levels_data: dict, tz_name: str, app_build: str,
    is_hindcast: bool = False,
) -> bytes:
    """Word report for app_beleid.py (the simplified policy/organiser
    view). Deliberately a SEPARATE entry point from generate_report_docx()
    rather than a thin wrapper around it: that function's per-level section
    hard-requires PYROX multi-day results (`pyrox_result["groups"]`), which
    this app never computes, by design (see its module docstring -- "NOT
    shown: ... multi-day PYROX context"). Reusing it here would either
    force app_beleid.py to compute data it deliberately hides on screen, or
    silently reintroduce exactly the raw/uncalibrated figures and
    PROVISIONAL caveats that app was built to leave out. This function
    instead mirrors app_beleid.py's own screen: day overview first, then
    one zoomed race-window section per level, each ending at HESTIA's own
    post-finish simulation window rather than an arbitrary margin.

    levels_data : dict, level_label -> {
        "finish": Timestamp, "post_finish_end": Timestamp,
        "hestia_result": dict | None, "met_value": float | None,
        "duration_minutes": float | None,
    }
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)

    generated_at = pd.Timestamp.now(tz=tz_name)
    title = ("PYROX \u2014 Policy Report (HINDCAST, observed weather)"
             if is_hindcast else "PYROX \u2014 Policy Report")
    _add_title_section(doc, city_name, exp_start, generated_at, app_build,
                       report_title=title)
    if is_hindcast:
        note = doc.add_paragraph()
        note.add_run(
            "\U0001F4DC HINDCAST: this report uses observed historical "
            "weather (Open-Meteo archive) for a date that has already "
            "passed, not a forecast. Intended for checking this model's "
            "prediction against a real, known event -- e.g. comparing "
            "against an actual_* outcome logged in a prediction record for "
            "the same date."
        ).italic = True
        doc.add_paragraph()
    _add_day_overview_section(doc, weather_df, exp_date, exp_start)

    for level_label, d in levels_data.items():
        _add_race_window_section(
            doc, level_label, weather_df, exp_start, d["finish"],
            d.get("post_finish_end", d["finish"]), d.get("hestia_result"),
            d.get("met_value"), d.get("duration_minutes"),
        )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        "This report presents model outputs and findings only, matching "
        "what the PYROX Beleid view showed on screen -- it does not include "
        "recommendations for operational measures. For full methodology, "
        "raw/uncalibrated figures, and multi-day PYROX context, use the "
        "PYROX Participants view."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_report_docx(
    city_name: str, exp_start, weather_df: pd.DataFrame,
    per_level_exposure: dict, pyrox_result: dict, label_for: dict,
    hestia_results: dict, finish, flag_warnings: list, hestia_warnings: list,
    tz_name: str, app_build: str, met_by_level_label: dict = None,
    duration_by_level_label: dict = None,
) -> bytes:
    """Builds the report and returns it as bytes, ready for a Streamlit
    download_button. No file is written to disk."""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)

    generated_at = pd.Timestamp.now(tz=tz_name)
    _add_title_section(doc, city_name, exp_start, generated_at, app_build)
    _add_weather_section(doc, weather_df, exp_start, finish)
    _add_per_level_section(doc, per_level_exposure, pyrox_result, label_for,
                           hestia_results, pd.Timestamp(exp_start).date(),
                           exp_start, finish, met_by_level_label or {},
                           duration_by_level_label or {})
    _add_divergence_section(doc, flag_warnings, hestia_warnings)
    _add_limitations_section(doc, hestia_results)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
